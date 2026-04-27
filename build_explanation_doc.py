"""Generate Stock_Project_Detailed_Explanation.docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent

def style_doc(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x11, 0x33, 0x66)

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic

def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

doc = Document()
style_doc(doc)

add_title(doc, "Stock Price Direction Predictor")
add_para(doc, "A detailed, step-by-step explanation of every concept, module, "
         "and line of code in the project.", italic=True)
doc.add_paragraph()

# =====================================================================
doc.add_heading("Part 1 — Project Overview", level=1)
# =====================================================================
add_para(doc,
    "This project predicts whether a stock's next-day closing price will go UP or DOWN. "
    "Unlike the Customer Churn project, which used a static CSV downloaded once from Kaggle, "
    "this project fetches live data from Yahoo Finance every time you run it. That single "
    "change — from static file to live API — introduces several new engineering and ML concepts.")

add_para(doc, "What makes this project different:", bold=True)
add_bullets(doc, [
    "Live data via API: every run pulls the latest prices from the internet.",
    "Technical indicator features: we engineer ~22 new columns from raw price data using financial formulas.",
    "Time-based splitting: we split data by date, not randomly, to prevent the model from seeing the future.",
    "TimeSeriesSplit for cross-validation: respects temporal order during hyperparameter tuning.",
    "Modular code: 4 separate Python files (modules) instead of everything in one notebook.",
    "On-demand retraining: the Streamlit app can retrain the model with one click.",
    "Backtesting: the app shows how the model would have performed historically.",
])

add_para(doc, "Project architecture:", bold=True)
add_code(doc,
    "data_fetcher.py  →  feature_engineer.py  →  trainer.py\n"
    "       ↓                     ↓                   ↓\n"
    "   Live OHLCV           22 features          Saved model\n"
    "       ↓                     ↓                   ↓\n"
    "       └──────────  app.py imports all three  ───┘")

# =====================================================================
doc.add_heading("Part 2 — What is an API?", level=1)
# =====================================================================
add_para(doc,
    "API stands for Application Programming Interface. In simple terms, it is a way for your code "
    "to talk to another computer over the internet and ask for data. When you use yfinance, here is "
    "what happens behind the scenes:")
add_bullets(doc, [
    "Your code calls yf.download('AAPL', period='2y', interval='1d').",
    "The yfinance library sends an HTTP request to Yahoo Finance's servers: 'Give me Apple's daily prices for the last 2 years.'",
    "Yahoo's server looks up the data in its database.",
    "The server sends the data back as structured text (JSON or CSV format).",
    "yfinance parses that text and converts it into a pandas DataFrame.",
    "You receive a DataFrame with columns: Date, Open, High, Low, Close, Volume.",
])
add_para(doc,
    "The key advantage over a static CSV: every time you call the API, you get today's latest data. "
    "The CSV you downloaded for the churn project was frozen the moment you downloaded it — it never "
    "updated. An API call is live.")

add_para(doc,
    "The yfinance library hides all the HTTP complexity. Without it, you would need to construct URLs, "
    "handle authentication, parse JSON, handle errors, and manage rate limits. Libraries like yfinance "
    "wrap all of that into a single function call.")

# =====================================================================
doc.add_heading("Part 3 — Module 1: data_fetcher.py (line-by-line)", level=1)
# =====================================================================
add_para(doc, "This module is responsible for all communication with the Yahoo Finance API.", bold=True)

add_code(doc, 'import yfinance as yf')
add_para(doc,
    "This imports the yfinance library and gives it the short alias 'yf' so we can write yf.download() "
    "instead of yfinance.download().")

doc.add_heading("Function 1: fetch_stock_data", level=2)
add_code(doc,
    'def fetch_stock_data(ticker="AAPL", period="2y", interval="1d"):\n'
    '    data = yf.download(ticker, period=period, interval=interval)\n'
    '    data = _flatten_columns(data)\n'
    '    data = data.dropna()\n'
    '    data = data.reset_index()\n'
    '    return data')
add_para(doc, "Line-by-line explanation:", bold=True)
add_bullets(doc, [
    "def fetch_stock_data(ticker='AAPL', ...) — defines a function with default parameters. If you call fetch_stock_data() with no arguments, it defaults to Apple, 2 years, daily data. But you can override any parameter: fetch_stock_data('TSLA', '1y').",
    "yf.download(ticker, ...) — the actual API call. It sends a request to Yahoo Finance and returns a DataFrame. The ticker parameter tells it which stock, period tells it how far back, interval tells it the candle size.",
    "_flatten_columns(data) — newer versions of yfinance return MultiIndex columns like ('Close', 'AAPL'). This helper function flattens them to just 'Close', 'Open', etc., which is cleaner to work with.",
    "data.dropna() — removes any rows with missing values. Some trading days might have gaps (holidays, data errors). Dropping them prevents NaN values from propagating through calculations.",
    "data.reset_index() — the DataFrame comes with Date as the index (row label). reset_index() turns it into a regular column, which is easier to work with for plotting and feature engineering.",
])

doc.add_heading("Function 2: fetch_multiple_stocks", level=2)
add_code(doc,
    'def fetch_multiple_stocks(tickers, period="2y", interval="1d"):\n'
    '    results = {}\n'
    '    for ticker in tickers:\n'
    '        results[ticker] = fetch_stock_data(ticker, period, interval)\n'
    '    return results')
add_para(doc, "Explanation:", bold=True)
add_bullets(doc, [
    "results = {} — creates an empty Python dictionary. A dictionary maps keys to values, like a real dictionary maps words to definitions.",
    "for ticker in tickers: — loops through each ticker in the list. If tickers = ['AAPL', 'MSFT'], the loop runs twice: once with ticker = 'AAPL', once with ticker = 'MSFT'.",
    "results[ticker] = fetch_stock_data(...) — stores the returned DataFrame in the dictionary under the ticker name as the key. After the loop, results looks like {'AAPL': <DataFrame>, 'MSFT': <DataFrame>}.",
    "return results — returns the dictionary AFTER the loop finishes. A common beginner mistake is putting return inside the loop, which makes it return after the first ticker only.",
])

doc.add_heading("Function 3: get_latest_price", level=2)
add_para(doc,
    "This fetches only the last 5 days and returns the most recent day's data as a dictionary. "
    "It uses data.iloc[-1] — iloc means 'integer location', and -1 means 'the last row'. "
    "This is useful when the app just needs today's price without downloading 2 years of history.")

doc.add_heading("The if __name__ == '__main__' block", level=2)
add_code(doc,
    'if __name__ == "__main__":\n'
    '    df = fetch_stock_data("AAPL")\n'
    '    print(df.head())')
add_para(doc,
    "This is a Python pattern you will see in every module. It means: 'only run this code if I am "
    "executing this file directly (python data_fetcher.py), NOT when another file imports me.' "
    "When app.py does 'from data_fetcher import fetch_stock_data', the __name__ variable is set to "
    "'data_fetcher' (not '__main__'), so the test code is skipped. This lets each module be both "
    "a library (importable) and a standalone script (testable).")

# =====================================================================
doc.add_heading("Part 4 — Module 2: feature_engineer.py (line-by-line)", level=1)
# =====================================================================
add_para(doc,
    "This module transforms raw OHLCV data into features a machine learning model can learn from. "
    "The raw data only has 5 columns (Open, High, Low, Close, Volume). A model cannot learn much from "
    "those alone — if AAPL closed at $150 today, that number means nothing without context. But 'the price "
    "is above its 50-day average, RSI is 72 (overbought), and volume spiked' — now there is a pattern.", bold=True)

doc.add_heading("What is OHLCV?", level=2)
add_bullets(doc, [
    "Open — the price at the start of the trading day.",
    "High — the highest price during the day.",
    "Low — the lowest price during the day.",
    "Close — the price at the end of the day (the most important one for predictions).",
    "Volume — how many shares were traded that day.",
])

doc.add_heading("Price Features", level=2)
add_code(doc,
    "data['returns_1d'] = close.pct_change(1)\n"
    "data['returns_5d'] = close.pct_change(5)\n"
    "data['log_return'] = np.log(close / close.shift(1))")
add_bullets(doc, [
    "pct_change(1) calculates the percentage change from yesterday to today: (today - yesterday) / yesterday. If the stock went from $100 to $102, the return is 0.02 (2%).",
    "pct_change(5) does the same but over 5 days — it captures the weekly trend.",
    "log_return is the logarithmic return. Finance uses log returns because they are additive (you can sum daily log returns to get the total return), whereas regular returns are multiplicative.",
    "close.shift(1) shifts the entire column down by 1 row — so each row now sees yesterday's close. This is how we create 'lagged' features.",
])

doc.add_heading("Moving Averages", level=2)
add_code(doc,
    "data['sma_5'] = close.rolling(5).mean()\n"
    "data['sma_20'] = close.rolling(20).mean()\n"
    "data['ema_12'] = close.ewm(span=12).mean()\n"
    "data['sma_cross'] = (data['sma_5'] > data['sma_20']).astype(int)")
add_bullets(doc, [
    "SMA (Simple Moving Average): the average close price over the last N days. SMA_5 is a fast signal (reacts quickly), SMA_50 is a slow signal (smooths out noise).",
    "rolling(5).mean() — creates a sliding window of 5 rows and takes the mean of each window. On day 10, it averages days 6-10. On day 11, it averages days 7-11. And so on.",
    "EMA (Exponential Moving Average): like SMA but gives more weight to recent days. It reacts faster to price changes. ewm(span=12) means the exponential window has a span of 12 days.",
    "sma_cross: 1 if the fast average (5-day) is above the slow average (20-day), else 0. When the fast crosses above the slow, traders call it a 'golden cross' (bullish). Below is a 'death cross' (bearish).",
])

doc.add_heading("Momentum Indicators (ta library)", level=2)
add_code(doc,
    "data['rsi_14'] = ta.momentum.RSIIndicator(close, window=14).rsi()\n"
    "macd_obj = ta.trend.MACD(close)\n"
    "data['macd'] = macd_obj.macd()")
add_bullets(doc, [
    "RSI (Relative Strength Index): measures how fast and how much the price has been going up or down. Ranges from 0 to 100. Above 70 = 'overbought' (price may fall soon). Below 30 = 'oversold' (price may rise soon). Window of 14 days is the industry standard.",
    "MACD (Moving Average Convergence Divergence): the difference between the 12-day EMA and the 26-day EMA. When MACD crosses above its signal line, it is a bullish signal. The 'ta' library calculates all of this for us — we just call the function.",
    "Stochastic Oscillator (%K): compares today's close to the high-low range over a period. High values mean the close is near the top of its range (potentially overbought).",
])

doc.add_heading("Volatility Indicators", level=2)
add_code(doc,
    "bb = ta.volatility.BollingerBands(close)\n"
    "data['bb_high'] = bb.bollinger_hband()\n"
    "data['bb_low'] = bb.bollinger_lband()\n"
    "data['atr_14'] = ta.volatility.AverageTrueRange(high, low, close).average_true_range()")
add_bullets(doc, [
    "Bollinger Bands: a band drawn 2 standard deviations above and below the 20-day SMA. When the price touches the upper band, it is 'stretched' high. When it touches the lower band, it is 'stretched' low. The width of the band measures volatility.",
    "ATR (Average True Range): measures how much the price moves on an average day. High ATR = volatile stock. Low ATR = calm stock. Useful because a $2 move on a $200 stock is very different from a $2 move on a $20 stock.",
])

doc.add_heading("Volume Indicators", level=2)
add_code(doc,
    "data['volume_sma_20'] = volume.rolling(20).mean()\n"
    "data['volume_ratio'] = volume / data['volume_sma_20']\n"
    "data['obv'] = ta.volume.OnBalanceVolumeIndicator(close, volume).on_balance_volume()")
add_bullets(doc, [
    "volume_ratio: today's volume divided by the 20-day average volume. A ratio of 2.0 means twice the normal volume — something significant is happening. High volume confirms price moves; low volume suggests the move might not last.",
    "OBV (On-Balance Volume): a running total that adds volume on up days and subtracts volume on down days. Rising OBV with rising price confirms the trend. Rising OBV with falling price suggests a reversal is coming.",
])

doc.add_heading("The Target Variable", level=2)
add_code(doc, "data['target'] = (close.shift(-1) > close).astype(int)")
add_para(doc,
    "This is the most important line. shift(-1) looks at TOMORROW's close price. If tomorrow's close "
    "is higher than today's, target = 1 (UP). Otherwise target = 0 (DOWN). This is what the model learns "
    "to predict. We drop the last row because we do not know tomorrow's price for it yet.")

# =====================================================================
doc.add_heading("Part 5 — Module 3: trainer.py (line-by-line)", level=1)
# =====================================================================

doc.add_heading("Time-Based Splitting (the most important concept)", level=2)
add_code(doc,
    "split_idx = int(len(df) * (1 - test_ratio))\n"
    "train = df.iloc[:split_idx]\n"
    "test  = df.iloc[split_idx:]")
add_para(doc,
    "In the churn project, we used random splitting (train_test_split with shuffle). That is fine for "
    "customer data because customers are independent of each other. But stock prices are sequential — "
    "today's price depends on yesterday's. If we randomly shuffled, a data point from 2026 could end up "
    "in training and a point from 2025 in testing. The model would literally use the future to predict "
    "the past — this is called 'data leakage' or 'look-ahead bias'.", bold=True)
add_para(doc,
    "The fix: a TIME-BASED split. Everything before a cutoff date is training, everything after is testing. "
    "This mimics how the model would actually be used in real life — it trains on past data and predicts "
    "the future.")
add_code(doc,
    "TRAIN: Jan 2024 ──────────────── Sep 2025\n"
    "TEST:                            Sep 2025 ── Apr 2026")

doc.add_heading("TimeSeriesSplit for Cross-Validation", level=2)
add_code(doc, "cv=TimeSeriesSplit(n_splits=5)")
add_para(doc,
    "In the churn project we used StratifiedKFold, which randomly shuffles data into 5 folds. Here we use "
    "TimeSeriesSplit instead, which creates expanding training windows:")
add_code(doc,
    "Fold 1: Train [---]     Test [--]\n"
    "Fold 2: Train [------]  Test [--]\n"
    "Fold 3: Train [---------] Test [--]\n"
    "Fold 4: Train [------------] Test [--]\n"
    "Fold 5: Train [---------------] Test [--]")
add_para(doc,
    "Each fold uses all previous data for training and the next chunk for testing. The test set always "
    "comes AFTER the training set in time. This is the correct way to validate time-series models.")

doc.add_heading("The Three Models", level=2)
add_para(doc, "We train the same three models as in the churn project, for the same reasons:")
add_bullets(doc, [
    "Logistic Regression — linear baseline. Fast, interpretable, but limited to linear patterns.",
    "Random Forest — bagging ensemble of 200 trees. Captures non-linear patterns.",
    "XGBoost — boosting ensemble. Usually the best performer on tabular data. Each tree corrects the errors of the previous one.",
])

doc.add_heading("GridSearchCV Tuning", level=2)
add_para(doc,
    "We tune 5 hyperparameters of XGBoost across 108 combinations × 5 time-series folds = 540 model fits. "
    "The hyperparameters control the bias-variance tradeoff:")
add_bullets(doc, [
    "max_depth [3, 5, 7] — deeper trees fit more complex patterns but risk overfitting.",
    "learning_rate [0.01, 0.05, 0.1] — smaller values make the model learn slower but more carefully.",
    "n_estimators [100, 200, 300] — more trees capture more patterns but take longer and risk overfitting.",
    "subsample [0.8, 1.0] — fraction of rows used per tree. Less than 1.0 adds randomness (regularisation).",
    "colsample_bytree [0.8, 1.0] — fraction of features used per tree. Same idea: adds randomness.",
])

doc.add_heading("The Retrain Pipeline", level=2)
add_code(doc,
    "def retrain_pipeline(ticker='AAPL', period='2y'):\n"
    "    df = fetch_stock_data(ticker, period)\n"
    "    data, feature_cols = engineer_features(df)\n"
    "    X_train, X_test, ... = prepare_data(data, feature_cols)\n"
    "    results = train_and_evaluate(...)\n"
    "    best_model = tune_best_model(...)\n"
    "    save_artifacts(best_model, scaler, feature_names)\n"
    "    return ...")
add_para(doc,
    "This function chains every step into a single call. The Streamlit app's 'Retrain' button calls this "
    "function. It fetches fresh data, re-engineers features, retrains all models, tunes the best one, and "
    "saves new artifacts. One function call = a complete model refresh with the latest market data.")

# =====================================================================
doc.add_heading("Part 6 — Module 4: app.py (the Streamlit Dashboard)", level=1)
# =====================================================================

doc.add_heading("How Streamlit Works", level=2)
add_para(doc,
    "Streamlit turns a Python script into a web app. Every time the user interacts with a widget (clicks "
    "a button, moves a slider), Streamlit re-runs the entire script from top to bottom. This is different "
    "from Flask or Django where you define routes and handlers.")

doc.add_heading("Key Streamlit Concepts Used", level=2)
add_bullets(doc, [
    "st.set_page_config() — sets the browser tab title and icon. Must be the first Streamlit command.",
    "st.sidebar — creates inputs in the left sidebar. Keeps the main area clean for results.",
    "st.columns(3) — creates 3 side-by-side columns for metrics.",
    "st.metric() — displays a number with a label, like a KPI dashboard.",
    "st.plotly_chart() — renders an interactive Plotly chart.",
    "st.pyplot() — renders a Matplotlib figure (used for SHAP plots).",
    "st.spinner() — shows a loading animation while code runs.",
    "st.expander() — creates a collapsible section.",
    "@st.cache_resource — caches the model and SHAP explainer so they are only loaded once, not on every re-run. Without this, the app would reload the model every time you click anything.",
])

doc.add_heading("The Prediction Flow", level=2)
add_para(doc, "When the user clicks 'Predict', here is what happens step by step:")
add_bullets(doc, [
    "1. fetch_stock_data(ticker) — calls the Yahoo Finance API to get live data.",
    "2. engineer_features(df) — creates 22 technical indicator features from the raw data.",
    "3. Take the LAST row of the engineered data — this is 'today', the most recent complete trading day.",
    "4. scaler.transform() — scales the features using the same scaler that was fitted during training.",
    "5. reindex(columns=feature_names) — ensures the columns are in the exact same order as during training.",
    "6. model.predict_proba() — returns the probability of UP (class 1). If > 0.5, predict UP; else DOWN.",
    "7. SHAP explainer — breaks down that single prediction into feature contributions.",
    "8. Display everything: prediction, confidence, SHAP waterfall, backtest chart.",
])

doc.add_heading("The Backtest Chart", level=2)
add_para(doc,
    "The backtest shows two lines: 'Buy & Hold' (you bought the stock and held it through the test period) "
    "vs 'Model Strategy' (you only hold the stock on days the model predicted UP, and sell/stay flat on "
    "DOWN days). If the model strategy line is above the Buy & Hold line, the model adds value. "
    "A disclaimer is added because past performance does not guarantee future results.")

# =====================================================================
doc.add_heading("Part 7 — The Notebook (How It Ties Everything Together)", level=1)
# =====================================================================
add_para(doc,
    "The notebook is a documented research walkthrough. It imports from the 3 modules and produces "
    "visualisations, model artifacts, and a summary. Here is what each cell does:")

add_bullets(doc, [
    "Cell 1 (Imports): loads all libraries and imports our custom modules using sys.path.insert(0, '../src'). This tells Python where to find our files.",
    "Cell 2 (Data Fetching): calls the API and prints the timestamp — proving this is live, not static.",
    "Cell 3 (Price Visualisation): 3 plots — price history with moving averages, volume bars, and daily returns distribution.",
    "Cell 4 (Feature Engineering): calls engineer_features() and shows the correlation heatmap — which features correlate most with tomorrow's direction.",
    "Cell 5 (Technical Indicators): a 3-panel chart showing Bollinger Bands, RSI, and MACD — the three most important technical indicators.",
    "Cell 6 (Data Preparation): time-based split with explicit date ranges printed to prove no leakage.",
    "Cell 7 (Model Training): trains all 3 models and produces comparison bar chart + ROC curves.",
    "Cell 8 (Hyperparameter Tuning): GridSearchCV with TimeSeriesSplit — 540 fits.",
    "Cell 9 (Evaluation): confusion matrix + precision-recall curve for the tuned model.",
    "Cell 10 (SHAP): 3 plots — bar importance, beeswarm, and waterfall for the highest-confidence prediction.",
    "Cell 11 (Save): saves model, scaler, and feature names to the models/ folder.",
    "Cell 12 (Live Prediction): fetches today's data and makes a real prediction — the demo moment.",
    "Cell 13 (Summary): prints everything in one block — metrics, top features, artifact counts, timestamp.",
])

# =====================================================================
doc.add_heading("Part 8 — Key Differences from the Churn Project", level=1)
# =====================================================================
add_para(doc, "Here is a comparison table of every major difference:", bold=True)

table = doc.add_table(rows=10, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Aspect", "Churn Project", "Stock Project"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

rows = [
    ("Data source", "Static Kaggle CSV", "Live Yahoo Finance API"),
    ("Features", "Raw dataset columns", "22 engineered technical indicators"),
    ("Train/test split", "Random stratified", "Time-based (no future leakage)"),
    ("Cross-validation", "StratifiedKFold", "TimeSeriesSplit"),
    ("Code structure", "Everything in one notebook", "4 separate importable modules"),
    ("App inputs", "User enters customer profile", "App fetches live data automatically"),
    ("Retraining", "Not possible", "One-click retrain in the app"),
    ("Evaluation extra", "None", "Backtest chart: model vs buy-and-hold"),
    ("Domain knowledge", "Binary yes/no features", "Finance: RSI, MACD, Bollinger, OBV"),
]
for i, (a, b, c) in enumerate(rows, 1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = c

# =====================================================================
doc.add_heading("Part 9 — Concepts You Can Now Explain in Interviews", level=1)
# =====================================================================
add_bullets(doc, [
    "What is an API and how does yfinance work under the hood.",
    "OHLCV data and why Close is the most important column for predictions.",
    "Technical indicators: RSI, MACD, Bollinger Bands, moving averages, OBV — what each measures and why traders use them.",
    "Feature engineering: transforming raw data into signals a model can learn from.",
    "The shift() function: creating lagged features and target variables in time-series.",
    "Time-based splitting vs random splitting, and why random splitting causes data leakage in time-series.",
    "TimeSeriesSplit: the expanding-window cross-validation method for temporal data.",
    "Modular code: separating concerns into importable modules vs monolithic notebooks.",
    "@st.cache_resource: why caching matters for app performance.",
    "Backtesting: simulating how a trading strategy would have performed historically.",
    "reindex(): ensuring column alignment between training and inference — critical for deployment.",
    "The retrain pattern: one-call pipelines that fetch, engineer, train, tune, and save.",
])

# =====================================================================
doc.add_heading("Part 10 — How to Present This Project", level=1)
# =====================================================================
add_para(doc,
    "When presenting, emphasise the PROGRESSION from the churn project to this one. You went from "
    "static data to live APIs, from random splits to time-based splits, from a single notebook to "
    "modular architecture, and from a form-based app to one that fetches its own data. Each step "
    "reflects a real-world production concern.")

add_para(doc, "Key sentences to include:", bold=True)
add_bullets(doc, [
    "'Unlike my previous project, this one works with live data — every prediction uses today's actual market data fetched via the Yahoo Finance API.'",
    "'I used a time-based train/test split to prevent look-ahead bias, which is critical in financial ML.'",
    "'The code is modular — four separate files that the notebook and the app both import from, which mirrors production code structure.'",
    "'The app includes a retrain button so the model can be updated with the latest data without touching the code.'",
    "'SHAP explanations show which technical indicators drove each prediction, making the model interpretable for stakeholders.'",
])

path = OUT / "Stock_Project_Detailed_Explanation.docx"
doc.save(path)
print("Saved:", path)
